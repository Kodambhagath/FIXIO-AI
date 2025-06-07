import os
import subprocess
import logging
from docx import Document
import pandas as pd

def convert_document(input_path, output_path, target_format):
    """
    Convert a document to a different format.
    
    Args:
        input_path (str): Path to input document file
        output_path (str): Path to save the converted document
        target_format (str): Target format (e.g., 'pdf', 'docx', 'txt')
    
    Returns:
        tuple: (success, message)
    """
    # Get input file extension
    input_ext = os.path.splitext(input_path)[1].lower().replace('.', '')
    
    try:
        # Word document (docx/doc) conversions
        if input_ext in ['docx', 'doc']:
            if target_format == 'pdf':
                return convert_word_to_pdf(input_path, output_path)
            elif target_format == 'txt':
                return convert_word_to_txt(input_path, output_path)
        
        # Excel document (xlsx/xls) conversions
        elif input_ext in ['xlsx', 'xls']:
            if target_format == 'pdf':
                return convert_excel_to_pdf(input_path, output_path)
            elif target_format == 'csv':
                return convert_excel_to_csv(input_path, output_path)
        
        # Text file conversions
        elif input_ext == 'txt':
            if target_format == 'docx':
                return convert_txt_to_docx(input_path, output_path)
            elif target_format == 'pdf':
                # First convert to docx, then to PDF
                temp_docx = output_path.replace('.pdf', '.docx')
                success, message = convert_txt_to_docx(input_path, temp_docx)
                if success:
                    return convert_word_to_pdf(temp_docx, output_path)
                else:
                    return False, message
        
        # Unsupported conversion
        return False, "This conversion is not supported"
        
    except Exception as e:
        logging.error(f"Error converting document: {str(e)}")
        return False, str(e)

def convert_word_to_pdf(input_path, output_path):
    """
    Convert a Word document to PDF.
    Since we cannot use libraries that would require additional system dependencies like
    LibreOffice or wkhtmltopdf, we'll create a simple PDF with PyPDF2.
    
    This is a simplified version - in a production environment, you'd use a more robust solution.
    
    Args:
        input_path (str): Path to input Word document
        output_path (str): Path to save the PDF
    
    Returns:
        tuple: (success, message)
    """
    try:
        import PyPDF2
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        # Extract text from Word document
        doc = Document(input_path)
        text = '\n\n'.join([para.text for para in doc.paragraphs])
        
        # Create a PDF with the text
        c = canvas.Canvas(output_path, pagesize=letter)
        width, height = letter
        
        # Split text into lines
        lines = text.split('\n')
        y = height - 50  # Start near the top
        
        for line in lines:
            if y < 50:  # If we're near the bottom, create a new page
                c.showPage()
                y = height - 50
            
            c.drawString(50, y, line)
            y -= 15  # Move down for the next line
        
        c.save()
        
        return True, "Conversion successful"
    
    except Exception as e:
        logging.error(f"Error converting Word to PDF: {str(e)}")
        return False, "Error converting Word to PDF"

def convert_word_to_txt(input_path, output_path):
    """
    Convert a Word document to a plain text file.
    
    Args:
        input_path (str): Path to input Word document
        output_path (str): Path to save the text file
    
    Returns:
        tuple: (success, message)
    """
    try:
        doc = Document(input_path)
        
        with open(output_path, 'w', encoding='utf-8') as f:
            for para in doc.paragraphs:
                f.write(para.text + '\n')
        
        return True, "Conversion successful"
    
    except Exception as e:
        logging.error(f"Error converting Word to text: {str(e)}")
        return False, "Error converting Word to text"

def convert_excel_to_pdf(input_path, output_path):
    """
    Convert an Excel file to PDF.
    
    Args:
        input_path (str): Path to input Excel file
        output_path (str): Path to save the PDF
    
    Returns:
        tuple: (success, message)
    """
    try:
        import PyPDF2
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import Table, TableStyle
        from reportlab.lib import colors
        
        # Read Excel file
        df = pd.read_excel(input_path)
        
        # Create a PDF with the data
        c = canvas.Canvas(output_path, pagesize=letter)
        width, height = letter
        
        # Convert DataFrame to a list of lists for the table
        data = [df.columns.tolist()] + df.values.tolist()
        
        # Create a table with the data
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        # Draw the table on the PDF
        table.wrapOn(c, width - 100, height)
        table.drawOn(c, 50, height - 100)
        
        c.save()
        
        return True, "Conversion successful"
    
    except Exception as e:
        logging.error(f"Error converting Excel to PDF: {str(e)}")
        return False, "Error converting Excel to PDF"

def convert_excel_to_csv(input_path, output_path):
    """
    Convert an Excel file to CSV.
    
    Args:
        input_path (str): Path to input Excel file
        output_path (str): Path to save the CSV file
    
    Returns:
        tuple: (success, message)
    """
    try:
        df = pd.read_excel(input_path)
        df.to_csv(output_path, index=False)
        
        return True, "Conversion successful"
    
    except Exception as e:
        logging.error(f"Error converting Excel to CSV: {str(e)}")
        return False, "Error converting Excel to CSV"

def convert_txt_to_docx(input_path, output_path):
    """
    Convert a text file to a Word document.
    
    Args:
        input_path (str): Path to input text file
        output_path (str): Path to save the Word document
    
    Returns:
        tuple: (success, message)
    """
    try:
        doc = Document()
        
        with open(input_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
            
            for line in lines:
                doc.add_paragraph(line.strip())
        
        doc.save(output_path)
        
        return True, "Conversion successful"
    
    except Exception as e:
        logging.error(f"Error converting text to Word: {str(e)}")
        return False, "Error converting text to Word"
